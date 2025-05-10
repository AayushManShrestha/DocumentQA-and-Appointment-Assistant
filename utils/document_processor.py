from langchain_core.documents import Document
from langchain_community.vectorstores import FAISS
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from PyPDF2 import PdfReader
from docx import Document as DocxReader
from io import BytesIO

def load_pdf(content):
    reader = PdfReader(BytesIO(content))
    return [Document(page_content=page.extract_text()) for page in reader.pages if page.extract_text()]

def load_txt(content):
    text = content.decode("utf-8")
    return [Document(page_content=text)]

def load_docx(content):
    doc = DocxReader(BytesIO(content))
    text = "\n".join([p.text for p in doc.paragraphs if p.text])
    return [Document(page_content=text)]

def process_documents(uploaded_files):
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    all_docs = []

    for file in uploaded_files:
        content = file.getvalue()
        ext = file.name.lower().split(".")[-1]

        if ext == "pdf":
            docs = load_pdf(content)
        elif ext == "txt":
            docs = load_txt(content)
        elif ext == "docx":
            docs = load_docx(content)
        else:
            continue

        all_docs.extend(docs)

    chunks = splitter.split_documents(all_docs)
    return FAISS.from_documents(chunks, embeddings)
